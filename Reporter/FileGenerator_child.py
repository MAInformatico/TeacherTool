from  TeacherTool.B2 import B2
from colorama import Fore
from docx import Document
import docxedit
import numpy as np
import os
import shutil
from Reporter.FileGenerator import FileGenerator

class filegenerator_child(FileGenerator):

    def __init__(self, language_level=None, students=None, num_exam=None, list_values=None):
        self.language_level = language_level
        self.students = students
        self.num_exam = num_exam
        self.values = list_values

    def get_info(self,average_overall,average_reading,average_listening,average_speaking,average_writing,average_use_english):
            language_level, students, num_exam = self.language_level, self.students, self.num_exam

            dict_results = {
                'Name' : [],
                'Reading': [],
                'Listening': [],
                'Writing': [],
                'Speaking': [],
                'Use_English': [],
                'Overall': []             
            }

            for st in range(0,students):
                average_test = []
                dict_results['Name'].append(self.values['Name'][int(st)])                
                
                instance = B2(dict_results['Name'],num_exam)                

                reading = instance.filter(str(self.values['Reading'][int(st)]))
                result = instance.percent_category(int(reading[0]),int(reading[1]))
                average_test.append(result)
                dict_results['Reading'].append(result)
                
                listening = instance.filter(str(self.values['Listening'][st]))
                result = instance.percent_category(int(listening[0]),int(listening[1]))
                average_test.append(result)
                dict_results['Listening'].append(result)

                writing = instance.filter(str(self.values['Writing'][st]))
                result = instance.percent_category(int(writing[0]),int(writing[1]))
                average_test.append(result)
                dict_results['Writing'].append(result)                

                speaking = instance.filter(str(self.values['Speaking'][st]))
                result = instance.percent_category(int(speaking[0]),int(speaking[1]))
                average_test.append(result)
                dict_results['Speaking'].append(result)                

                use_english = instance.filter(str(self.values['Use_English'][st]))
                result = instance.percent_category(int(speaking[0]),int(use_english[1]))
                average_test.append(result)
                dict_results['Use_English'].append(result)                 

                '''
                instance = B2(dict_results['Name'],num_exam)
                avg_reading = []
                avg_listening = []
                avg_writing = []
                avg_speaking = []
                avg_use_english = []
                for _ in range(0,num_exam):            
                    average_test = []
                    #Reading
                    aux = dict_results['Name'][students]
                    reading = instance.filter(str(aux))
                    result = instance.percent_category(int(reading[0]),int(reading[1]))
                    average_test.append(result)
                    avg_reading.append(result)
                    #Listening
                    aux = dict_results['Listening'][students]
                    listening = instance.filter(aux)
                    result = instance.percent_category(listening[0],listening[1])
                    average_test.append(result)
                    avg_listening.append(result)
                    #Writting
                    aux = dict_results['Writing'][students]
                    writing = instance.filter(aux)
                    result = instance.percent_category(writing[0],writing[1])
                    average_test.append(result)
                    avg_writing.append(result)
                    dict_results['Writing'].append(str(aux))
                    #Speaking
                    aux = dict_results['Speaking'][students]
                    speaking = instance.filter(aux)
                    result = instance.percent_category(speaking[0],speaking[1])
                    average_test.append(result)
                    avg_speaking.append(result)
                    dict_results['Speaking'].append(str(aux))
                    if (language_level[0] == "B2" and language_level[1] == 2) or (language_level[0] == "C1"):
                        #English use
                        aux = dict_results['Use_English'][students]
                        english_use = instance.filter(aux)                    
                        result = instance.percent_category(english_use[0],english_use[1])
                        average_test.append(result)
                        avg_use_english.append(result)
                        dict_results['Use_English'].append(str(aux))
                    else: #It is a B1 exam
                        dict_results['Use_English'].append(str('X'))
                    #calculate average and insert in Overall array
                    aux = np.mean(average_test)
                    dict_results['Overall'].append(np.ceil(aux))                
                average_overall.append(aux)
                
                average_reading.append(np.mean(avg_reading))
                average_listening.append(np.mean(avg_listening))
                average_writing.append(np.mean(avg_writing))
                average_speaking.append(np.mean(avg_speaking))
                if (language_level[0] == "B2" and language_level[1] == 2) or (language_level[0] == "C1"):
                    average_use_english.append(np.mean(avg_use_english))
                '''
                print(average_test)
                aux = np.mean(average_test)
                print(type(aux))
                dict_results['Overall'].append(np.ceil(aux))
                
            return language_level, dict_results, num_exam, students



    def generate_files(self,dict_notes,num_exam):
        valuetoclean = "np.float64("        
        name_file = 'test.docx'
        print(dict_notes)
        if num_exam == 1 and len(dict_notes['Name']) == 1:
            document = Document(name_file)
            docxedit.add_text_in_table(document.tables[0], row_num=1, column_num=0, new_string=str(dict_notes['Name']).replace('[\'','').replace('\']',''))
            docxedit.add_text_in_table(document.tables[0], row_num=1, column_num=1, new_string=str(dict_notes['Reading']).replace('[\'','').replace('\']',''))
            docxedit.add_text_in_table(document.tables[0], row_num=1, column_num=2, new_string=str(dict_notes['Use_English']).replace('[\'','').replace('\']',''))
            docxedit.add_text_in_table(document.tables[0], row_num=1, column_num=3, new_string=str(dict_notes['Writing']).replace('[\'','').replace('\']',''))
            docxedit.add_text_in_table(document.tables[0], row_num=1, column_num=4, new_string=str(dict_notes['Listening']).replace('[\'','').replace('\']',''))
            docxedit.add_text_in_table(document.tables[0], row_num=1, column_num=5, new_string=str(dict_notes['Speaking']).replace('[\'','').replace('\']',''))
            docxedit.add_text_in_table(document.tables[0], row_num=1, column_num=6, new_string=str(dict_notes['Overall']).replace(valuetoclean,'').replace(')','').replace('[','').replace(']',''))
            document.save(name_file)
        elif num_exam == 1 and len(dict_notes['Name']) > 1:
            document = Document(name_file)
            for i in range(0,len(dict_notes['Name'])):            
                docxedit.add_text_in_table(document.tables[0], row_num=i+1, column_num=0, new_string=str(dict_notes['Name'][i]).replace('[\'','').replace('\']',''))
                docxedit.add_text_in_table(document.tables[0], row_num=i+1, column_num=1, new_string=str(dict_notes['Reading'][i]).replace('[\'','').replace('\']',''))
                docxedit.add_text_in_table(document.tables[0], row_num=i+1, column_num=2, new_string=str(dict_notes['Use_English'][i]).replace('[\'','').replace('\']',''))
                docxedit.add_text_in_table(document.tables[0], row_num=i+1, column_num=3, new_string=str(dict_notes['Writing'][i]).replace('[\'','').replace('\']',''))
                docxedit.add_text_in_table(document.tables[0], row_num=i+1, column_num=4, new_string=str(dict_notes['Listening'][i]).replace('[\'','').replace('\']',''))
                docxedit.add_text_in_table(document.tables[0], row_num=i+1, column_num=5, new_string=str(dict_notes['Speaking'][i]).replace('[\'','').replace('\']',''))
                docxedit.add_text_in_table(document.tables[0], row_num=i+1, column_num=6, new_string=str(dict_notes['Overall'][int(i)]).replace(valuetoclean,'').replace(')','').replace('[','').replace(']',''))
            document.save(name_file)
        file_oldname = os.path.join(str(os.getcwd()), name_file)
        file_newname_newfile = os.path.join(str(os.getcwd()), "test_results" + ".docx")
        os.rename(file_oldname, file_newname_newfile)
        shutil.copy(file_newname_newfile, os.path.join(str(os.getcwd()), "test.docx"))
        return file_newname_newfile        